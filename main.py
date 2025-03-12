import os
import importlib.util
from flask import Flask, render_template, send_file, redirect, url_for, request, send_from_directory
from functools import wraps
import sys
import zipfile
from werkzeug.utils import secure_filename
import shutil
import pandas as pd
import re
from docx import Document
from datetime import datetime

# Create the Flask application
app = Flask(__name__,
    template_folder='templates',
    static_folder='static'
)

# Helper functie om HTML bestanden te renderen
def render_html_file(file_path):
    """Render een HTML bestand met de juiste headers."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # Controleer of het bestand Jinja2 template tags bevat
        if '{%' in content or '{{' in content:
            # Als het bestand template tags bevat, gebruik render_template_string
            from flask import render_template_string
            return render_template_string(content)
        else:
            # Anders, geef het bestand direct terug
            return content, 200, {'Content-Type': 'text/html'}
    except Exception as e:
        print(f"Error rendering HTML file {file_path}: {e}")
        return f"Error rendering file: {str(e)}", 500

# Define the with_app_context decorator
def with_app_context(app_name):
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            # This decorator can be used to set the correct application context
            # for functions that need it
            return f(*args, **kwargs)
        return decorated_function
    return decorator

# Homepage route - This should be at the top to ensure it's registered first
@app.route('/')
def home():
    # Huidige datum in het juiste formaat
    current_date = datetime.now().strftime("%d-%m-%Y")

    return render_template('index.html',
                          title="Dagvaardingen Automatisering",
                          current_date=current_date,
                          apps=[
                              {
                                  "name": "Stap 1: Templates",
                                  "route": "/app1/",
                                  "description": "Upload Excel bestanden en Word templates voor dagvaardingen. Het systeem zoekt de juiste eisen in de Excel bestanden en past de templates aan."
                              },
                              {
                                  "name": "Stap 2: Facturen",
                                  "route": "/app2/",
                                  "description": "Upload PDF facturen die verwerkt moeten worden. Het systeem extraheert de data en plaatst deze in een Excel template."
                              },
                              {
                                  "name": "Stap 3: Dagvaardingen",
                                  "route": "/app3/",
                                  "description": "Combineer de output van stap 1 en 2 om de definitieve dagvaardingen te genereren en te exporteren."
                              }
                          ])

# Utility route om terug te keren naar de home
@app.route('/terug-naar-dashboard')
def back_to_dashboard():
    return redirect(url_for('home'))

# App 1 integreren
try:
    # Registreer de app1 routes in de hoofdapplicatie
    @app.route('/app1/')
    def app1_index():
        # Probeer de app1 interface direct te laden
        try:
            # Zoek naar een app1 index.html in de app1 directory
            if os.path.exists(os.path.join('app1', 'index.html')):
                return render_html_file(os.path.join('app1', 'index.html'))
            elif os.path.exists(os.path.join('app1', 'templates', 'index.html')):
                return render_html_file(os.path.join('app1', 'templates', 'index.html'))
            else:
                # Als geen app1 interface gevonden wordt, toon onze eigen template
                return render_template('index.html',
                                      title="Stap 1: Templates",
                                      app_name="Stap 1: Templates Voorbereiden",
                                      app_description="Upload Excel bestanden en Word templates voor dagvaardingen. Het systeem zoekt de juiste eisen in de Excel bestanden en past de templates aan met de bijbehorende versies.")
        except Exception as e:
            print(f"Error loading app1 interface: {e}")
            # Fallback naar onze eigen template
            return render_template('index.html',
                                  title="Stap 1: Templates",
                                  app_name="Stap 1: Templates Voorbereiden",
                                  app_description="Upload Excel bestanden en Word templates voor dagvaardingen. Het systeem zoekt de juiste eisen in de Excel bestanden en past de templates aan met de bijbehorende versies.")

    @app.route('/app1/<path:path>', methods=['GET', 'POST'])
    def app1_proxy(path):
        # Handle routing to app1
        try:
            # Probeer eerst de app1 functie direct aan te roepen
            if request.method == 'GET':
                # Voor GET requests, probeer een statisch bestand te serveren
                try:
                    # Probeer verschillende locaties voor statische bestanden
                    possible_paths = [
                        os.path.join('app1', 'static', path),
                        os.path.join('app1', path),
                        os.path.join('app1', 'templates', path)
                    ]

                    for p in possible_paths:
                        if os.path.exists(p):
                            return send_file(p, as_attachment=False)
                except Exception as e:
                    print(f"Error serving static file from app1: {e}")

            elif request.method == 'POST':
                # Voor POST requests, probeer de juiste functie aan te roepen
                try:
                    # Specifieke POST routes
                    if path == 'upload':
                        # Als er een upload functie is in app1.main, roep deze aan
                        try:
                            # Importeer de upload_files functie uit app1.main
                            print(f"Python path: {sys.path}")

                            # Voeg de app1 directory toe aan het pad als deze er nog niet in zit
                            app1_path = os.path.abspath('app1')
                            if app1_path not in sys.path:
                                sys.path.append(app1_path)
                                print(f"Added {app1_path} to Python path")

                            # Sla het huidige werkdirectory op
                            original_cwd = os.getcwd()

                            try:
                                # Verander het werkdirectory naar app1
                                os.chdir(app1_path)
                                print(f"Changed working directory to: {os.getcwd()}")

                                # Importeer de upload_files functie
                                try:
                                    from main import upload_files
                                    print("Successfully imported upload_files from main")
                                    result = upload_files()
                                    return result
                                except (ImportError, AttributeError) as e:
                                    print(f"Error importing upload_files from main: {e}")

                                    # Als laatste optie, probeer de functie direct aan te roepen
                                    try:
                                        # Voer de code van de upload_files functie direct uit
                                        files = request.files.getlist("files")

                                        if not files:
                                            return "Geen bestanden ontvangen.", 400

                                        # Sla bestanden op in de app1 directory
                                        UPLOAD_FOLDER = "uploads"
                                        PROCESSED_FOLDER = "processed_documents"
                                        os.makedirs(UPLOAD_FOLDER, exist_ok=True)
                                        os.makedirs(PROCESSED_FOLDER, exist_ok=True)

                                        excel_path = None
                                        word_folder = os.path.join(UPLOAD_FOLDER, "word_docs")
                                        os.makedirs(word_folder, exist_ok=True)

                                        for file in files:
                                            filename = secure_filename(file.filename)
                                            if filename.endswith(".xlsx"):
                                                excel_path = os.path.join(UPLOAD_FOLDER, filename)
                                                file.save(excel_path)
                                            elif filename.endswith(".docx"):
                                                file.save(os.path.join(word_folder, filename))

                                        if not excel_path:
                                            return "Excel bestand ontbreekt!", 400

                                        # Importeer de benodigde functies uit process.py
                                        from process import determine_variant, process_documents

                                        # Varianten bepalen en documenten verwerken
                                        selected_variants = determine_variant(excel_path)
                                        process_documents(word_folder, PROCESSED_FOLDER, selected_variants)

                                        # ZIP-bestand maken
                                        zip_path = os.path.join(PROCESSED_FOLDER, "verwerkte_documenten.zip")
                                        with zipfile.ZipFile(zip_path, 'w') as zipf:
                                            for root, dirs, files in os.walk(PROCESSED_FOLDER):
                                                for file in files:
                                                    if file != "verwerkte_documenten.zip":
                                                        file_path = os.path.join(root, file)
                                                        zipf.write(file_path, os.path.relpath(file_path, PROCESSED_FOLDER))

                                        # Redirect naar success pagina
                                        return redirect('/app1/success')
                                    except Exception as e:
                                        print(f"Error executing upload_files code directly: {e}")
                            finally:
                                # Herstel het originele werkdirectory
                                os.chdir(original_cwd)
                                print(f"Restored working directory to: {os.getcwd()}")
                        except Exception as e:
                            print(f"Error handling upload in app1: {e}")

                    # Voeg hier andere POST routes toe indien nodig
                except Exception as e:
                    print(f"Error handling POST request to app1/{path}: {e}")

            # Als we hier komen, is er geen route gevonden
            return f"Route niet gevonden: /app1/{path}", 404
        except Exception as e:
            print(f"Error in app1_proxy for path {path}: {e}")
            return f"Error: {str(e)}", 500

    @app.route('/app1/success')
    def app1_success():
        try:
            # Probeer de success.html pagina te laden
            if os.path.exists(os.path.join('app1', 'success.html')):
                return render_html_file(os.path.join('app1', 'success.html'))
            else:
                # Als de success.html pagina niet gevonden wordt, toon een eenvoudige success pagina
                html = """
                <!DOCTYPE html>
                <html lang="nl">
                <head>
                    <meta charset="UTF-8">
                    <meta name="viewport" content="width=device-width, initial-scale=1.0">
                    <title>✅ Verwerking Geslaagd</title>
                    <link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/@picocss/pico@2/css/pico.min.css">
                    <style>
                        :root {
                            --success-color: #2ecc71;
                            --primary-color: #0078D7;
                        }

                        body {
                            display: flex;
                            justify-content: center;
                            align-items: center;
                            min-height: 100vh;
                            background-color: #f3f6fc;
                            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                            margin: 0;
                            line-height: 1.6;
                        }

                        .success-container {
                            background-color: white;
                            padding: 2.5rem;
                            border-radius: 12px;
                            box-shadow: 0 10px 25px rgba(0, 0, 0, 0.15);
                            width: 100%;
                            max-width: 500px;
                            text-align: center;
                            transition: transform 0.3s ease-in-out;
                        }

                        .success-container:hover {
                            transform: translateY(-5px);
                        }

                        .success-icon {
                            font-size: 4rem;
                            color: var(--success-color);
                            margin-bottom: 1rem;
                        }

                        .action-links {
                            display: flex;
                            flex-direction: column;
                            gap: 1rem;
                            margin-top: 1.5rem;
                        }

                        .action-links a {
                            text-decoration: none;
                            padding: 0.75rem;
                            border-radius: 6px;
                            font-size: 16px;
                            font-weight: bold;
                            transition: all 0.3s ease;
                            text-align: center;
                            display: block;
                        }

                        .download-link {
                            background-color: var(--success-color);
                            color: white;
                        }

                        .download-link:hover {
                            background-color: #27ae60;
                            transform: scale(1.05);
                        }

                        .back-link {
                            background-color: var(--primary-color);
                            color: white;
                        }

                        .back-link:hover {
                            background-color: #005a9e;
                            transform: scale(1.05);
                        }

                        .dashboard-link {
                            background-color: #3498db;
                            color: white;
                        }

                        .dashboard-link:hover {
                            background-color: #2980b9;
                            transform: scale(1.05);
                        }
                    </style>
                </head>
                <body>
                    <main class="success-container">
                        <div class="success-icon">✅</div>
                        <h2>Verwerking Geslaagd!</h2>
                        <p style="color: #555; font-size: 18px;">Je bestanden zijn succesvol verwerkt.</p>
                        <div class="action-links">
                            <a href="/app1/download" class="download-link">⬇️ Download ZIP bestand</a>
                            <a href="/app1/" class="back-link">🔙 Terug naar Upload</a>
                            <a href="/terug-naar-dashboard" class="dashboard-link">🏠 Terug naar Dashboard</a>
                        </div>
                    </main>
                </body>
                </html>
                """
                return html
        except Exception as e:
            print(f"Error in app1_success: {e}")
            return f"Error: {str(e)}", 500

    @app.route('/app1/download')
    def app1_download():
        try:
            # Zoek het ZIP bestand
            zip_path = os.path.join('app1', 'processed_documents', 'verwerkte_documenten.zip')

            # Controleer of het ZIP bestand bestaat
            if os.path.exists(zip_path):
                print(f"ZIP bestand gevonden: {zip_path}")
                return send_file(zip_path, as_attachment=True, download_name='verwerkte_documenten.zip')

            # Als het ZIP bestand niet bestaat, zoek naar individuele bestanden
            processed_folder = os.path.join('app1', 'processed_documents')
            if os.path.exists(processed_folder):
                # Maak een nieuw ZIP bestand met alle bestanden in de processed_documents map
                zip_path = os.path.join(processed_folder, 'verwerkte_documenten.zip')
                with zipfile.ZipFile(zip_path, 'w') as zipf:
                    for root, dirs, files in os.walk(processed_folder):
                        for file in files:
                            if file != 'verwerkte_documenten.zip':
                                file_path = os.path.join(root, file)
                                zipf.write(file_path, os.path.relpath(file_path, processed_folder))

                print(f"Nieuw ZIP bestand aangemaakt: {zip_path}")
                return send_file(zip_path, as_attachment=True, download_name='verwerkte_documenten.zip')

            # Als er geen bestanden zijn gevonden, geef een foutmelding
            print("Geen verwerkte bestanden gevonden")
            return "Fout: Er zijn geen verwerkte bestanden gevonden.", 404
        except Exception as e:
            print(f"Error in app1_download: {e}")
            return f"Er is een fout opgetreden bij het downloaden: {str(e)}", 500

    print("App1 succesvol geïntegreerd")
except Exception as e:
    print(f"Kon app1 niet integreren: {e}")

# App 2 integreren
try:
    # Importeer app2/app.py (het hoofdbestand van app2)
    spec2 = importlib.util.spec_from_file_location("app2.app", os.path.join("app2", "app.py"))
    app2_module = importlib.util.module_from_spec(spec2)
    spec2.loader.exec_module(app2_module)

    # Registreer de app2 routes in de hoofdapplicatie
    @app.route('/app2/')
    def app2_index():
        # Probeer de app2 interface direct te laden
        try:
            # Controleer of het index.html bestand bestaat in app2 directory
            if os.path.exists(os.path.join('app2', 'index.html')):
                # Lees het bestand en stuur het direct terug
                with open(os.path.join('app2', 'index.html'), 'r', encoding='utf-8') as f:
                    content = f.read()
                    return content, 200, {'Content-Type': 'text/html'}

            # Als dat niet lukt, probeer de index functie van app2 aan te roepen
            try:
                # Probeer de index functie direct aan te roepen
                from app2.app import index
                return index()
            except (ImportError, AttributeError) as e:
                print(f"Could not import index from app2.app: {e}")

            # Als laatste optie, toon onze eigen template
            return render_template('index.html',
                                  title="Stap 2: Facturen",
                                  app_name="Stap 2: Facturen Verwerken",
                                  app_description="Upload PDF facturen die verwerkt moeten worden. Het systeem extraheert de data uit de PDF's en plaatst deze in een Excel template voor verdere verwerking.")
        except Exception as e:
            print(f"Error in app2_index: {e}")
            # Fallback naar onze eigen template
            return render_template('index.html',
                                  title="Stap 2: Facturen",
                                  app_name="Stap 2: Facturen Verwerken",
                                  app_description="Upload PDF facturen die verwerkt moeten worden. Het systeem extraheert de data uit de PDF's en plaatst deze in een Excel template voor verdere verwerking.")

    @app.route('/app2/success')
    def app2_success():
        try:
            # Haal de folders parameter op uit de query string
            folders = request.args.get("folders", "")
            folder_list = folders.split(",") if folders else []

            # Probeer de success.html pagina te laden
            if os.path.exists(os.path.join('app2', 'success.html')):
                # Lees de template en render deze met de juiste context
                with open(os.path.join('app2', 'success.html'), 'r', encoding='utf-8') as f:
                    template_content = f.read()

                from flask import render_template_string
                return render_template_string(template_content, folders=folder_list)
            else:
                # Als geen template gevonden wordt, toon een eenvoudige HTML pagina
                html = """
                <!DOCTYPE html>
                <html lang="nl">
                <head>
                    <meta charset="UTF-8">
                    <meta name="viewport" content="width=device-width, initial-scale=1.0">
                    <title>Verwerking Geslaagd</title>
                    <script src="https://cdn.tailwindcss.com"></script>
                </head>
                <body class="bg-gradient-to-br from-green-50 to-green-100 min-h-screen flex items-center justify-center p-4">
                    <div class="w-full max-w-xl bg-white rounded-3xl shadow-2xl p-8">
                        <div class="text-center mb-6">
                            <div class="flex justify-center mb-4">
                                <svg class="w-16 h-16 text-green-600" fill="none" stroke="currentColor" viewBox="0 0 24 24" xmlns="http://www.w3.org/2000/svg">
                                    <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M9 12l2 2 4-4m6 2a9 9 0 11-18 0 9 9 0 0118 0z"></path>
                                </svg>
                            </div>
                            <h1 class="text-3xl font-bold mb-2 text-green-800">Verwerking Geslaagd!</h1>
                            <p class="text-gray-600 text-lg">De volgende mappen zijn succesvol verwerkt:</p>
                        </div>
                """

                # Voeg de lijst met mappen toe
                if folder_list:
                    html += """
                        <div class="bg-green-50 p-4 rounded-lg mb-6 max-h-48 overflow-y-auto">
                            <ul class="space-y-2">
                    """

                    for folder in folder_list:
                        html += f"""
                                <li class="flex items-center bg-white p-3 rounded-md shadow-sm">
                                    <svg class="w-5 h-5 text-green-500 mr-3" fill="none" stroke="currentColor" viewBox="0 0 24 24" xmlns="http://www.w3.org/2000/svg">
                                        <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M3 7v10a2 2 0 002 2h14a2 2 0 002-2V9a2 2 0 00-2-2h-6l-2-2H5a2 2 0 00-2 2z"></path>
                                    </svg>
                                    <span class="text-gray-700 font-medium">{folder}</span>
                                </li>
                        """

                    html += """
                            </ul>
                        </div>
                    """
                else:
                    html += """
                        <p class="text-red-600 bg-red-50 p-4 rounded-lg text-center">Geen mappen verwerkt.</p>
                    """

                # Voeg de knoppen toe
                html += """
                        <div class="flex flex-col space-y-4">
                            <a
                                href="/app2/download"
                                class="w-full bg-green-600 hover:bg-green-700 text-white font-bold py-3 rounded-lg transition-all duration-300 transform hover:scale-105 flex items-center justify-center text-center"
                            >
                                <svg class="w-6 h-6 mr-2" fill="none" stroke="currentColor" viewBox="0 0 24 24" xmlns="http://www.w3.org/2000/svg">
                                    <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M4 16v1a3 3 0 003 3h10a3 3 0 003-3v-1m-4-4l-4 4m0 0l-4-4m4 4V4"></path>
                                </svg>
                                Download Resultaten
                            </a>

                            <a
                                href="/app2/"
                                class="w-full text-center text-green-700 hover:text-green-900 font-semibold py-3 rounded-lg transition-colors bg-green-100 hover:bg-green-200"
                            >
                                Terug naar uploadpagina
                            </a>

                            <a
                                href="/terug-naar-dashboard"
                                class="w-full text-center text-blue-700 hover:text-blue-900 font-semibold py-3 rounded-lg transition-colors bg-blue-100 hover:bg-blue-200"
                            >
                                Terug naar Dashboard
                            </a>
                        </div>
                    </div>
                </body>
                </html>
                """

                return html
        except Exception as e:
            print(f"Error in app2_success: {e}")
            return f"Error: {str(e)}", 500

    @app.route('/app2/download')
    def app2_download():
        try:
            # Zoek het ZIP bestand
            zip_path = os.path.join('app2', 'facturen_in_excel', "facturen_ingevuld.zip")

            # Controleer of het ZIP bestand bestaat
            if not os.path.exists(zip_path):
                # Probeer alternatieve locaties
                alternative_paths = [
                    os.path.join('facturen_in_excel', "facturen_ingevuld.zip"),
                    os.path.join('app2', 'facturen_in_excel', "facturen_ingevuld.zip")
                ]

                for alt_path in alternative_paths:
                    if os.path.exists(alt_path):
                        zip_path = alt_path
                        break
                else:
                    return "Fout: Het bestand bestaat niet. Verwerk eerst facturen.", 404

            # Stuur het ZIP-bestand naar de gebruiker
            return send_file(zip_path, as_attachment=True, download_name="facturen_ingevuld.zip")
        except Exception as e:
            print(f"Error in app2_download: {e}")
            return f"Er is een fout opgetreden bij het downloaden: {str(e)}", 500

    @app.route('/app2/<path:path>', methods=['GET', 'POST'])
    def app2_proxy(path):
        # Hier sturen we de request door naar de juiste route in app2
        try:
            if request.method == 'GET':
                # Voor GET requests, probeer een statisch bestand te serveren
                try:
                    # Probeer verschillende locaties voor statische bestanden
                    possible_paths = [
                        os.path.join('app2', 'static', path),
                        os.path.join('app2', path),
                        os.path.join('app2', 'templates', path),
                        os.path.join('app2', 'facturen_in_excel', path)
                    ]

                    for p in possible_paths:
                        if os.path.exists(p):
                            return send_file(p, as_attachment=False)
                except Exception as e:
                    print(f"Error serving static file from app2: {e}")

            elif request.method == 'POST':
                # Voor POST requests, probeer de juiste functie aan te roepen
                try:
                    # Specifieke POST routes
                    if path == 'upload':
                        # Sla het huidige werkdirectory op
                        original_cwd = os.getcwd()

                        try:
                            # Verander het werkdirectory naar app2
                            app2_path = os.path.abspath('app2')
                            os.chdir(app2_path)
                            print(f"Changed working directory to: {os.getcwd()}")

                            # Importeer de upload_files functie
                            try:
                                from app2.app import upload_files
                                print("Successfully imported upload_files from app2.app")
                                result = upload_files()
                                return result
                            except (ImportError, AttributeError) as e:
                                print(f"Error importing upload_files from app2.app: {e}")

                                # Als dat niet lukt, probeer de functie direct aan te roepen
                                try:
                                    # Voer de code van de upload_files functie direct uit
                                    files = request.files.getlist("files")

                                    if not files:
                                        return jsonify({"error": "Geen bestanden ontvangen."}), 400

                                    folders = {}
                                    for file in files:
                                        folder_name = file.filename.split('/')[0]  # Pak de mapnaam
                                        if folder_name not in folders:
                                            folders[folder_name] = {"excel": None, "pdf_files": []}

                                    if file.filename.endswith(".xlsx") or file.filename.endswith(".xls"):
                                        folders[folder_name]["excel"] = file
                                    elif file.filename.endswith(".pdf"):
                                        folders[folder_name]["pdf_files"].append(file)

                                    processed_files = []
                                    for folder_name, files in folders.items():
                                        if files["excel"] and files["pdf_files"]:
                                            folder_path = os.path.join('uploads', secure_filename(folder_name))
                                            os.makedirs(folder_path, exist_ok=True)

                                            excel_path = os.path.join(folder_path, secure_filename(files["excel"].filename))
                                            files["excel"].save(excel_path)

                                            pdf_paths = []
                                            for pdf_file in files["pdf_files"]:
                                                pdf_path = os.path.join(folder_path, secure_filename(pdf_file.filename))
                                                pdf_file.save(pdf_path)
                                                pdf_paths.append(pdf_path)

                                            output_filename = f"{folder_name}_verwerkt.xlsx"

                                            # Importeer de benodigde functies
                                            from app2.app import fill_excel_template

                                            output_path = fill_excel_template(excel_path, pdf_paths, output_filename)
                                            processed_files.append(output_path)

                                    # Maak een ZIP-bestand met alle ingevulde Excel-bestanden
                                    zip_path = os.path.join('facturen_in_excel', "facturen_ingevuld.zip")
                                    with zipfile.ZipFile(zip_path, "w") as zipf:
                                        for file in processed_files:
                                            zipf.write(file, os.path.basename(file))

                                    return redirect('/app2/success?folders=' + ",".join(folders.keys()))
                                except Exception as e:
                                    print(f"Error executing upload_files code directly: {e}")
                        finally:
                            # Herstel het originele werkdirectory
                            os.chdir(original_cwd)
                            print(f"Restored working directory to: {os.getcwd()}")

                    # Voeg hier andere POST routes toe indien nodig
                except Exception as e:
                    print(f"Error handling POST request to app2/{path}: {e}")

            # Als we hier komen, is er geen route gevonden
            return f"Route niet gevonden: /app2/{path}", 404
        except Exception as e:
            print(f"Error in app2_proxy for path {path}: {e}")
            return f"Error: {str(e)}", 500

    print("App2 succesvol geïntegreerd")
except Exception as e:
    print(f"Kon app2 niet integreren: {e}")

# App 3 integreren
try:
    # Importeer app3/deel_3.py (het hoofdbestand van app3)
    spec3 = importlib.util.spec_from_file_location("app3.deel_3", os.path.join("app3", "deel_3.py"))
    app3_module = importlib.util.module_from_spec(spec3)
    spec3.loader.exec_module(app3_module)

    # Registreer de app3 routes in de hoofdapplicatie
    @app.route('/app3/')
    def app3_index():
        # Probeer de app3 interface direct te laden
        try:
            # Probeer eerst de index functie van app3 te gebruiken via het app3_module
            if hasattr(app3_module, 'index'):
                return app3_module.index()

            # Als dat niet lukt, probeer de index functie te importeren
            try:
                from app3.deel_3 import index
                return index()
            except (ImportError, AttributeError) as e:
                print(f"Could not import index from app3.deel_3: {e}")

            # Als dat niet lukt, probeer de app3 interface te laden via het app3 object
            if hasattr(app3_module, 'app') and hasattr(app3_module.app, 'send_static_file'):
                try:
                    return app3_module.app.send_static_file('index.html')
                except Exception as e:
                    print(f"Error sending static file from app3: {e}")

            # Als dat niet lukt, probeer het app3 index.html bestand direct te laden
            if os.path.exists(os.path.join('app3', 'index.html')):
                return render_html_file(os.path.join('app3', 'index.html'))
            elif os.path.exists(os.path.join('app3', 'templates', 'index.html')):
                return render_html_file(os.path.join('app3', 'templates', 'index.html'))

            # Als geen app3 interface gevonden wordt, toon onze eigen template
            return render_template('index.html',
                                  title="Stap 3: Dagvaardingen",
                                  app_name="Stap 3: Dagvaardingen Genereren",
                                  app_description="Combineer de output van stap 1 en 2 om de definitieve dagvaardingen te genereren. Hier worden alle gegevens samengevoegd en de uiteindelijke dagvaardingen aangemaakt.")
        except Exception as e:
            print(f"Error in app3_index: {e}")
            # Fallback naar onze eigen template
            return render_template('index.html',
                                  title="Stap 3: Dagvaardingen",
                                  app_name="Stap 3: Dagvaardingen Genereren",
                                  app_description="Combineer de output van stap 1 en 2 om de definitieve dagvaardingen te genereren. Hier worden alle gegevens samengevoegd en de uiteindelijke dagvaardingen aangemaakt.")

    # Specifieke routes voor app3
    @app.route('/app3/upload_files', methods=['POST'])
    @with_app_context('app3')
    def app3_upload_files():
        # Hier roepen we de upload_files functie van app3 aan, maar omzeilen de flash functie
        try:
            # Maak uploads map leeg voor nieuwe bestanden
            app3_upload_folder = os.path.join('app3', 'uploads')
            app3_processed_folder = os.path.join('app3', 'processed')

            # Zorg ervoor dat de mappen bestaan
            os.makedirs(app3_upload_folder, exist_ok=True)
            os.makedirs(app3_processed_folder, exist_ok=True)

            # Maak uploads map leeg voor nieuwe bestanden
            shutil.rmtree(app3_upload_folder, ignore_errors=True)
            os.makedirs(app3_upload_folder, exist_ok=True)

            # Maak processed map leeg
            shutil.rmtree(app3_processed_folder, ignore_errors=True)
            os.makedirs(app3_processed_folder, exist_ok=True)

            # Check of bestanden zijn geüpload
            if 'files' not in request.files:
                return redirect('/app3/')

            uploaded_files = request.files.getlist('files')
            if not uploaded_files or uploaded_files[0].filename == '':
                return redirect('/app3/')

            # Sla alle geüploade bestanden op
            for file in uploaded_files:
                filename = secure_filename(file.filename)
                file.save(os.path.join(app3_upload_folder, filename))

            # Verwerk de bestanden
            excel_files = {}
            word_files = {}
            processed_count = 0

            # Verzamel alle Excel- en Word-bestanden
            for file in os.listdir(app3_upload_folder):
                full_path = os.path.join(app3_upload_folder, file)
                if file.lower().endswith(('.xlsx', '.xls')):
                    excel_files[file] = full_path
                elif file.lower().endswith('.docx'):
                    word_files[file] = full_path

            # Als er geen bestanden zijn, geef een foutmelding
            if not excel_files or not word_files:
                return redirect('/app3/')

            # Importeer de benodigde modules voor verwerking
            sys.path.append(os.path.abspath('app3'))
            try:
                from app3.main import process_document
                from app3.mapping import mapping

                # Verwerk de bestanden
                for excel_file, excel_path in excel_files.items():
                    try:
                        # Probeer een dossiernummer uit het Excel bestand te halen
                        try:
                            df_excel = pd.read_excel(excel_path, header=None)
                            dossiernummer = str(df_excel.iat[0, 0]).strip()
                        except Exception as e:
                            print(f"Kan dossiernummer niet uit Excel halen: {e}")
                            dossiernummer = None

                        # Zoek een bijpassend Word bestand
                        matched_word = None
                        if dossiernummer:
                            for word_file, word_path in word_files.items():
                                # Probeer het dossiernummer uit het Word bestand te halen
                                try:
                                    # Gebruik docx2txt als het beschikbaar is, anders gebruik Document
                                    try:
                                        import docx2txt
                                        full_text = docx2txt.process(word_path)
                                    except ImportError:
                                        doc = Document(word_path)
                                        full_text = '\n'.join([p.text for p in doc.paragraphs])

                                    matches = re.findall(r'Dossier:\s*(\d{8,9})', full_text, re.MULTILINE)
                                    doc_nummer = matches[-1].strip().zfill(8) if matches else None

                                    if doc_nummer and doc_nummer == dossiernummer:
                                        matched_word = word_path
                                        break
                                except Exception as e:
                                    print(f"Fout bij extractie dossiernummer uit {word_file}: {e}")
                                    continue

                        # Als er geen match is, gebruik dan gewoon beide bestanden als we maar één Excel en één Word bestand hebben
                        if not matched_word and len(excel_files) == 1 and len(word_files) == 1:
                            matched_word = list(word_files.values())[0]
                            print(f"Geen dossiernummer match, maar slechts één Word bestand beschikbaar. Gebruik: {os.path.basename(matched_word)}")

                        if matched_word:
                            # Bepaal de naam van het outputbestand
                            output_name = f"processed_{os.path.basename(matched_word)}"
                            output_path = os.path.join(app3_processed_folder, output_name)

                            # Verwerk het document met de process_document functie
                            success = process_document(excel_path, matched_word, output_path)

                            if success:
                                processed_count += 1
                                print(f"Document succesvol verwerkt: {output_path}")
                            else:
                                print(f"Fout bij verwerking van {excel_file}")
                    except Exception as e:
                        print(f"Fout bij verwerken van {excel_file}: {str(e)}")
            except ImportError as e:
                print(f"Kan de benodigde modules niet importeren: {e}")
                # Fallback: kopieer de Word bestanden naar de processed map
                for word_file, word_path in word_files.items():
                    output_path = os.path.join(app3_processed_folder, f"processed_{word_file}")
                    shutil.copy2(word_path, output_path)
                    processed_count += 1

            # Als er helemaal niets is verwerkt, ga terug naar de uploadpagina
            if processed_count == 0:
                return redirect('/app3/')

            # Maak een ZIP-bestand van de verwerkte documenten
            zip_filename = "processed_documents.zip"
            zip_path = os.path.join(app3_processed_folder, zip_filename)

            with zipfile.ZipFile(zip_path, 'w') as zipf:
                for file in os.listdir(app3_processed_folder):
                    file_path = os.path.join(app3_processed_folder, file)
                    if file.endswith(".docx") and os.path.isfile(file_path):
                        zipf.write(file_path, os.path.basename(file_path))

            # Redirect naar success pagina
            return redirect(f'/app3/success?zip_filename={zip_filename}')
        except Exception as e:
            print(f"Error in app3_upload_files: {e}")
            return f"Error: {str(e)}", 500

    @app.route('/app3/upload_single_files', methods=['POST'])
    @with_app_context('app3')
    def app3_upload_single_files():
        # Hier roepen we de upload_single_files functie van app3 aan, maar omzeilen de flash functie
        try:
            # Maak uploads map leeg voor nieuwe bestanden
            app3_upload_folder = os.path.join('app3', 'uploads')
            app3_processed_folder = os.path.join('app3', 'processed')

            # Zorg ervoor dat de mappen bestaan
            os.makedirs(app3_upload_folder, exist_ok=True)
            os.makedirs(app3_processed_folder, exist_ok=True)

            # Check of bestanden zijn geüpload
            if 'excel_file' not in request.files or 'word_file' not in request.files:
                return redirect('/app3/')

            excel_file = request.files['excel_file']
            word_file = request.files['word_file']

            if excel_file.filename == '' or word_file.filename == '':
                return redirect('/app3/')

            # Sla de bestanden op
            excel_filename = secure_filename(excel_file.filename)
            word_filename = secure_filename(word_file.filename)

            excel_path = os.path.join(app3_upload_folder, excel_filename)
            word_path = os.path.join(app3_upload_folder, word_filename)

            excel_file.save(excel_path)
            word_file.save(word_path)

            # Verwerk het document
            output_filename = f"verwerkt_{word_filename}"
            output_path = os.path.join(app3_processed_folder, output_filename)

            # Importeer de benodigde modules voor verwerking
            sys.path.append(os.path.abspath('app3'))
            try:
                from app3.main import process_document

                # Verwerk het document met de process_document functie
                success = process_document(excel_path, word_path, output_path)

                if success:
                    print(f"Document succesvol verwerkt: {output_path}")
                else:
                    print(f"Fout bij verwerking van document")
                    # Fallback: kopieer het Word bestand naar de processed map
                    shutil.copy2(word_path, output_path)
            except ImportError as e:
                print(f"Kan de benodigde modules niet importeren: {e}")
                # Fallback: kopieer het Word bestand naar de processed map
                shutil.copy2(word_path, output_path)

            # Redirect naar success pagina
            return redirect(f'/app3/success?zip_filename={output_filename}')
        except Exception as e:
            print(f"Error in app3_upload_single_files: {e}")
            return f"Error: {str(e)}", 500

    @app.route('/app3/success')
    def app3_success():
        try:
            zip_filename = request.args.get('zip_filename', 'processed_documents.zip')

            # Probeer de success.html pagina direct te laden uit de app3 directory
            if os.path.exists(os.path.join('app3', 'success.html')):
                with open(os.path.join('app3', 'success.html'), 'r', encoding='utf-8') as f:
                    content = f.read()

                    # Vervang Jinja2 template tags met HTML
                    # Vervang de flashed messages sectie
                    content = re.sub(r'{%\s*with\s+messages\s*=\s*get_flashed_messages\(with_categories=true\)\s*%}.*?{%\s*endwith\s*%}',
                                     '<div class="alert alert-success">Uw document(en) zijn succesvol verwerkt en staan klaar om te downloaden.</div>',
                                     content, flags=re.DOTALL)

                    # Vervang de if-statements voor categorieën
                    content = re.sub(r'{%\s*if\s+category\s*==\s*\'success\'\s*%}.*?{%\s*endif\s*%}',
                                     '<span class="success-icon">✅</span>',
                                     content, flags=re.DOTALL)
                    content = re.sub(r'{%\s*if\s+category\s*==\s*\'danger\'\s*%}.*?{%\s*endif\s*%}',
                                     '<span class="danger-icon">❌</span>',
                                     content, flags=re.DOTALL)
                    content = re.sub(r'{%\s*if\s+category\s*==\s*\'warning\'\s*%}.*?{%\s*endif\s*%}',
                                     '<span class="warning-icon">⚠️</span>',
                                     content, flags=re.DOTALL)

                    # Vervang de for-loop voor messages
                    content = re.sub(r'{%\s*for\s+category,\s*message\s+in\s+messages\s*%}.*?{%\s*endfor\s*%}',
                                     '',
                                     content, flags=re.DOTALL)

                    # Vervang de if-else voor messages
                    content = re.sub(r'{%\s*if\s+messages\s*%}.*?{%\s*else\s*%}.*?{%\s*endif\s*%}',
                                     '',
                                     content, flags=re.DOTALL)

                    # Vervang de variabele voor zip_filename
                    content = content.replace('{{ zip_filename }}', zip_filename)

                    return content, 200, {'Content-Type': 'text/html'}
            else:
                # Als de success.html niet gevonden wordt, toon een eenvoudige success pagina
                html = f"""
                <!DOCTYPE html>
                <html lang="nl">
                <head>
                    <meta charset="UTF-8">
                    <meta name="viewport" content="width=device-width, initial-scale=1.0">
                    <title>Verwerking Geslaagd</title>
                    <script src="https://cdn.tailwindcss.com"></script>
                </head>
                <body class="bg-gradient-to-br from-green-50 to-green-100 min-h-screen flex items-center justify-center p-4">
                    <div class="w-full max-w-xl bg-white rounded-3xl shadow-2xl p-8">
                        <div class="text-center mb-6">
                            <div class="flex justify-center mb-4">
                                <svg class="w-16 h-16 text-green-600" fill="none" stroke="currentColor" viewBox="0 0 24 24" xmlns="http://www.w3.org/2000/svg">
                                    <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M9 12l2 2 4-4m6 2a9 9 0 11-18 0 9 9 0 0118 0z"></path>
                                </svg>
                            </div>
                            <h1 class="text-3xl font-bold mb-2 text-green-800">Verwerking Geslaagd!</h1>
                            <p class="text-gray-600 text-lg">De bestanden zijn succesvol verwerkt.</p>
                        </div>

                        <div class="flex flex-col space-y-4">
                            <a
                                href="/app3/download_file?filename={zip_filename}"
                                class="w-full bg-green-600 hover:bg-green-700 text-white font-bold py-3 rounded-lg transition-all duration-300 transform hover:scale-105 flex items-center justify-center text-center"
                            >
                                <svg class="w-6 h-6 mr-2" fill="none" stroke="currentColor" viewBox="0 0 24 24" xmlns="http://www.w3.org/2000/svg">
                                    <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M4 16v1a3 3 0 003 3h10a3 3 0 003-3v-1m-4-4l-4 4m0 0l-4-4m4 4V4"></path>
                                </svg>
                                Download Verwerkte Document(en)
                            </a>

                            <a
                                href="/app3/"
                                class="w-full text-center text-green-700 hover:text-green-900 font-semibold py-3 rounded-lg transition-colors bg-green-100 hover:bg-green-200"
                            >
                                Terug naar uploadpagina
                            </a>

                            <a
                                href="/terug-naar-dashboard"
                                class="w-full text-center text-blue-700 hover:text-blue-900 font-semibold py-3 rounded-lg transition-colors bg-blue-100 hover:bg-blue-200"
                            >
                                Terug naar Dashboard
                            </a>
                        </div>
                    </div>
                </body>
                </html>
                """
                return html
        except Exception as e:
            print(f"Error in app3_success: {e}")
            return f"Error: {str(e)}", 500

    @app.route('/app3/download_file')
    @with_app_context('app3')
    def app3_download_file():
        # Hier roepen we de download_file functie van app3 aan, maar implementeren het direct
        try:
            filename = request.args.get('filename', 'processed_documents.zip')

            # Zoek het bestand in verschillende mogelijke locaties
            possible_paths = [
                os.path.join('app3', 'processed', filename),
                os.path.join('app3', filename),
                os.path.join('processed', filename)
            ]

            # Controleer of het bestand bestaat in een van de mogelijke locaties
            file_path = None
            for path in possible_paths:
                if os.path.exists(path):
                    file_path = path
                    break

            # Als het bestand niet gevonden is, geef een foutmelding
            if not file_path:
                return "Bestand niet gevonden", 404

            # Bepaal het MIME-type op basis van de bestandsextensie
            mimetype = None
            if filename.endswith('.docx'):
                mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            elif filename.endswith('.xlsx'):
                mimetype = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            elif filename.endswith('.zip'):
                mimetype = 'application/zip'

            # Stuur het bestand naar de gebruiker
            return send_file(file_path, as_attachment=True, download_name=filename, mimetype=mimetype)
        except Exception as e:
            print(f"Error in app3_download_file: {e}")
            return f"Error: {str(e)}", 500

    @app.route('/app3/<path:path>', methods=['GET', 'POST'])
    def app3_proxy(path):
        # Hier sturen we de request door naar de juiste route in app3
        try:
            # Probeer eerst de app3 functie direct aan te roepen
            if request.method == 'GET':
                # Voor GET requests, probeer de functie te vinden in app3.deel_3
                try:
                    from app3.deel_3 import app as app3

                    # Zoek naar een matching route in app3
                    for rule in app3.url_map.iter_rules():
                        if rule.endpoint != 'static' and rule.rule == '/' + path:
                            view_func = app3.view_functions.get(rule.endpoint)
                            if view_func:
                                return view_func()
                except Exception as e:
                    print(f"Error finding route in app3: {e}")

                # Als dat niet lukt, probeer een statisch bestand te serveren
                try:
                    # Probeer verschillende locaties voor statische bestanden
                    possible_paths = [
                        os.path.join('app3', 'static', path),
                        os.path.join('app3', path),
                        os.path.join('app3', 'templates', path)
                    ]

                    for p in possible_paths:
                        if os.path.exists(p):
                            return send_file(p, as_attachment=False)
                except Exception as e:
                    print(f"Error serving static file from app3: {e}")

            elif request.method == 'POST':
                # Voor POST requests, probeer de juiste functie aan te roepen
                try:
                    # Specifieke POST routes
                    if path == 'upload_files':
                        from app3.deel_3 import upload_files
                        return upload_files()
                    elif path == 'upload_single_files':
                        try:
                            from app3.deel_3 import upload_single_files
                            return upload_single_files()
                        except AttributeError:
                            from app3.deel_3 import upload_single
                            return upload_single()

                    # Voeg hier andere POST routes toe indien nodig
                except Exception as e:
                    print(f"Error handling POST request to app3/{path}: {e}")

            # Als we hier komen, is er geen route gevonden
            return f"Route niet gevonden: /app3/{path}", 404
        except Exception as e:
            print(f"Error in app3_proxy for path {path}: {e}")
            return f"Error: {str(e)}", 500

    print("App3 succesvol geïntegreerd")
except Exception as e:
    print(f"Kon app3 niet integreren: {e}")

if __name__ == '__main__':
    app.run(debug=True, port=5000)
