import os
import json
import shutil  
import pandas as pd
from docx import Document
import re
import docx2txt

def extract_dossiernummer(doc_path):
    """Extracts the correct dossier number from a Word document, ensuring it corresponds to the document being processed."""
    try:
        full_text = docx2txt.process(doc_path)

        # Debugging: Print de eerste 500 tekens van het document
        print(f"🔍 Geëxtraheerde tekst uit {doc_path}:\n{full_text[:500]}")

        # Zoek naar ALLE dossiernummers in het document
        matches = re.findall(r'Dossier:\s*(\d{8,9})', full_text, re.MULTILINE)

        if matches:
            # Kies het meest waarschijnlijke dossiernummer:
            dossiernummer = matches[-1].strip().zfill(8)  # Gebruik HET LAATSTE dossiernummer in de lijst
            print(f"✅ Correct dossiernummer geselecteerd: {dossiernummer} voor bestand {doc_path}")
            return dossiernummer
        else:
            print(f"⚠️ Geen Dossiernummer gevonden in {doc_path}")
    except Exception as e:
        print(f"❌ Fout bij extractie dossiernummer uit {doc_path}: {e}")

    return None

def determine_variant(excel_path):
    """Leest het Excel-bestand en bepaalt de variant voor elk dossiernummer."""
    df = pd.read_excel(excel_path, dtype=str).fillna("")
    
    variant_mapping = {
        "5015": "BPF_Bouw_5014/5015",
        "5016": "BPF_Bouw_5016",
        "5017": "BPF_Bouw_5017",
        "5019": "BPF_Bouw_5019",
        "5008": "BPF_Schoonmaak_5007/5008",
    }

    df.columns = df.columns.map(lambda x: str(x).strip())

    if "dossiernummer" not in df.columns:
        raise ValueError("❌ Kolom 'dossiernummer' ontbreekt in het Excel-bestand!")

    relevant_columns = {col: variant for col, variant in variant_mapping.items() if col in df.columns}

    if not relevant_columns:
        raise ValueError("❌ Geen verwachte variantkolommen gevonden in het Excel-bestand!")

    df["dossiernummer"] = df["dossiernummer"].astype(str).str.strip().apply(lambda x: x.zfill(8))

    selected_variants = {}

    for _, row in df.iterrows():
        dossiernummer = row["dossiernummer"]
        assigned_variant = None
        for col, variant in relevant_columns.items():
            if row.get(col, "").strip().upper() == "X":
                assigned_variant = variant
                print(f"✅ Dossiernummer {dossiernummer} toegewezen aan variant: {variant}")
                break  

        if assigned_variant:
            selected_variants[dossiernummer] = assigned_variant
        else:
            print(f"⚠️ Geen variant gevonden voor dossiernummer {dossiernummer}")

    if not selected_variants:
        print("⚠️ Geen enkel dossiernummer kreeg een variant toegewezen! Controleer je Excel-bestand.")
        print(f"🔎 Gekoppelde varianten uit Excel: {json.dumps(selected_variants, indent=2)}")

    return selected_variants

def clean_document(doc):
    """Reinigt het document zonder de originele structuur en opmaak te verliezen."""
    
    for para in doc.paragraphs:
        # Verwijdert achtergrondkleur en zet tekstkleur terug naar standaard
        for run in para.runs:
            run.font.highlight_color = None
            run.font.color.rgb = None
        
        # Verwijder alleen lege bulletpoints (lijsten zonder inhoud)
        if para.text.strip() in ["•", "-", "●", "*"]:
            para.text = ""

    # Verwijder extra lege regels maar behoud structuur
    for i in reversed(range(len(doc.paragraphs) - 1)):  
        if not doc.paragraphs[i].text.strip() and not doc.paragraphs[i + 1].text.strip():
            p = doc.paragraphs[i]._element
            p.getparent().remove(p)  # Verwijder de lege paragraaf veilig

    return doc

def process_bouw_document(doc, variant_data):
    """Past de bouw-specifieke wijzigingen toe op het document."""
    print("🔨 Bouwbewerking gestart...")

    posities = {
        "positie1": ("OP VERZOEK VAN", "GEDAGVAARD"),
        "positie2": ("FEITEN EN OMSTANDIGHEDEN", "Vordering"),
        "positie3": ("Gedaagde partij moet aan eisende partij een vordering betalen van € x1.", "Aanmelding werknemers"),
        "positie4": ("Eisende partij vordert, dat bij vonnis, uitvoerbaar bij voorraad, gedaagde partij wordt veroordeeld:", "tot betaling van de proceskosten")
    }
    
    for positie, (start_text, end_text) in posities.items():
        start_idx, end_idx = None, None
        for i, para in enumerate(doc.paragraphs):
            if start_text in para.text:
                start_idx = i + 1
            if end_text in para.text and start_idx is not None:
                end_idx = i
                break
        if start_idx is not None and end_idx is not None:
            for i in range(start_idx, end_idx):
                if doc.paragraphs[i].text.strip():
                    doc.paragraphs[i].text = ""

            if positie in variant_data:
                doc.paragraphs[start_idx].text = variant_data[positie]

    return clean_document(doc)

def process_schoonmaak_document(doc, variant_data):
    """Past de schoonmaak-specifieke wijzigingen toe op het document."""
    print("🧼 Schoonmaakbewerking gestart...")

    posities = {
        "positie1": ("OP VERZOEK VAN", "GEDAGVAARD"),
        "positie2": ("FEITEN EN OMSTANDIGHEDEN", "Vordering"),
        "positie3": ("Gedaagde partij moet aan eisende partij een vordering betalen van € x1.", "Aanmelding werknemers")
    }

    for positie, (start_text, end_text) in posities.items():
        start_idx, end_idx = None, None
        for i, para in enumerate(doc.paragraphs):
            if start_text in para.text:
                start_idx = i + 1
            if end_text in para.text and start_idx is not None:
                end_idx = i
                break
        if start_idx is not None and end_idx is not None:
            for i in range(start_idx, end_idx):
                if doc.paragraphs[i].text.strip():
                    doc.paragraphs[i].text = ""

            if positie in variant_data:
                doc.paragraphs[start_idx].text = variant_data[positie]

    return clean_document(doc)

def apply_correct_formatting(paragraph, text, bold_phrases, italic_phrases):
    """
    Ensures specific words are italicized or bolded correctly, including standalone words.
    """
    remaining_text = text  # Preserve original text
    current_pos = 0  

    # Check if the entire paragraph is just "Bijdrage" and apply italics directly
    if text.strip() in italic_phrases:
        paragraph.clear()
        run = paragraph.add_run(text.strip())
        run.italic = True
        return  # Exit function since it's already handled

    while current_pos < len(remaining_text):
        # Search for the next bold or italic phrase
        bold_match = next((phrase for phrase in bold_phrases if phrase in remaining_text[current_pos:]), None)
        italic_match = next((phrase for phrase in italic_phrases if phrase in remaining_text[current_pos:]), None)

        # Determine which phrase comes first
        next_match = None
        match_type = None

        if bold_match and italic_match:
            if remaining_text.index(bold_match) < remaining_text.index(italic_match):
                next_match = bold_match
                match_type = "bold"
            else:
                next_match = italic_match
                match_type = "italic"
        elif bold_match:
            next_match = bold_match
            match_type = "bold"
        elif italic_match:
            next_match = italic_match
            match_type = "italic"

        # If no match, add the remaining text as normal and stop
        if not next_match:
            paragraph.add_run(remaining_text[current_pos:])
            break

        # Add normal text before the match
        match_start = remaining_text.index(next_match, current_pos)
        paragraph.add_run(remaining_text[current_pos:match_start])

        # Add matched text with correct formatting
        run = paragraph.add_run(next_match)
        if match_type == "bold":
            run.bold = True
        elif match_type == "italic":
            run.italic = True

        # Move position forward
        current_pos = match_start + len(next_match)

bold_phrases = [
    "Stichting Bedrijfstakpensioenfonds voor de Bouwnijverheid",
    "Stichting Bedrijfstakpensioenfonds en voor de Bouwnijverheid",
    "Stichting Opleidings- en Ontwikkelingsfonds Afbouw",
    "Stichting Sociaal en Werkgelegenheidsfonds Timmerindustrie",
    "Stichting Opleidings- en Ontwikkelingsfonds Bouw & Infra",
    "Stichting Aanvullingsfonds Bouw & Infra",
    "Stichting Bedrijfstakpensioenfonds voor het Schoonmaak- en Glazenwassersbedrijf",
    "Stichting Raad voor Arbeidsverhoudingen Schoonmaak- en Glazenwassersbranche (Ras)",
    "GEDAGVAARD",
    "OP VERZOEK VAN",
    "MET AANZEGGING",
    "MET HET DOEL OM",
    "FEITEN EN OMSTANDIGHEDEN",
    "Aanmelding werknemers",
    "Nota's",
    "Verrichte werkzaamheden",
    "Vertragingsrente",
    "Buitengerechtelijke kosten",
    "Toerekening eventuele betalingen",
    "EVENTUEEL VERWEER EN DE WEERLEGGING DAARVAN",
    "BEWIJSLAST EN BEWIJSMIDDELEN",
    "OP WELKE GRONDEN",
    "MET HET DOEL OM",
    "FEITEN EN OMSTANDIGHEDEN",
    "Stichting Aanvullingsfonds Bouw & Infra"


]
italic_phrases = ["Premies", "Te late betaling", " Bijdrage", "Stichting Bedrijfstakpensioenfonds", "Stichting O&OA", "Stichting SWT", "Sociaal Fonds BIKUDAK", "Stichting OOB&I", "Stichting AB&I", "Subtotaal"]

def process_documents(input_folder, output_folder, selected_variants):
    """Verwerkt alle documenten en past de juiste functie toe (Bouw of Schoonmaak)."""
    # Gebruik het absolute pad naar het variants_v2.json bestand
    variants_file_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "variants_v2.json")
    
    try:
        with open(variants_file_path, "r", encoding="utf-8") as file:
            variants = json.load(file)
    except FileNotFoundError:
        print(f"⚠️ Bestand niet gevonden: {variants_file_path}")
        # Probeer alternatieve locaties
        alt_paths = [
            "variants_v2.json",
            os.path.join("app1", "variants_v2.json"),
            os.path.join(os.getcwd(), "app1", "variants_v2.json"),
            os.path.join(os.getcwd(), "variants_v2.json")
        ]
        
        for alt_path in alt_paths:
            try:
                print(f"🔍 Proberen alternatief pad: {alt_path}")
                with open(alt_path, "r", encoding="utf-8") as file:
                    variants = json.load(file)
                print(f"✅ Bestand gevonden op: {alt_path}")
                break
            except FileNotFoundError:
                continue
        else:
            print("❌ Kon variants_v2.json niet vinden op alle mogelijke locaties")
            variants = {}

    os.makedirs(output_folder, exist_ok=True)
    processed_files = []

    for filename in os.listdir(input_folder):
        if not filename.endswith(".docx"):
            continue

        doc_path = os.path.join(input_folder, filename)
        dossiernummer = extract_dossiernummer(doc_path)

        if not dossiernummer or dossiernummer.strip() == "" or dossiernummer not in selected_variants:
            print(f"⚠️ Ongeldig of ontbrekend dossiernummer in {filename}, overslaan...")
            continue

        # **Koppel de juiste variant per dossiernummer**
        variant = selected_variants.get(dossiernummer)
        
        # Debugging: Controleer de variant-koppeling
        print(f"🔍 Koppeling: Dossiernummer {dossiernummer} -> Variant {variant}")

        if not variant:
            print(f"⚠️ Geen variant gevonden voor dossiernummer {dossiernummer}, overslaan...")
            continue

        # **Laad het document opnieuw per iteratie om hergebruik te voorkomen**
        doc = Document(doc_path)

        if "Bouw" in variant:
            print(f"🔨 Bouw-document gedetecteerd: {filename}, verwerken als {variant}...")
            processed_doc = process_bouw_document(doc, variants.get(variant, {}))
        elif "Schoonmaak" in variant:
            print(f"🧹 Schoonmaak-document gedetecteerd: {filename}, verwerken als {variant}...")
            processed_doc = process_schoonmaak_document(doc, variants.get(variant, {}))
        else:
            print(f"⚠️ Onbekende variant voor {filename}: {variant}")
            continue

        # ✅ Pas de correcte opmaak toe op alle paragrafen
        for para in processed_doc.paragraphs:
            text = para.text
            para.clear()  # Verwijdert de originele tekst maar behoudt de paragraafstructuur
            apply_correct_formatting(para, text, bold_phrases, italic_phrases)

        # **Correcte output-bestandsnaam maken per dossiernummer**
        output_filename = f"processed_{dossiernummer}_{filename}"
        output_file_path = os.path.join(output_folder, output_filename)

        processed_doc.save(output_file_path)
        processed_files.append(output_file_path)

        # ✅ Verwijder het originele bestand na verwerking
        if os.path.exists(doc_path):
            os.remove(doc_path)
            print(f"🗑️ Origineel bestand verwijderd: {doc_path}")
        else:
            print(f"⚠️ Kon bestand niet verwijderen, bestaat niet: {doc_path}")

    return f"✅ {len(processed_files)} document(en) verwerkt en opgeslagen in: {output_folder}"