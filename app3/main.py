import pandas as pd
import re
from docx import Document
import os
import sys
from mapping import mapping
from placeholders import all_placeholders

def haal_excel_waarde(df, excel_cel):
    try:
        col_letter = ''.join(filter(str.isalpha, excel_cel))
        row_number = int(''.join(filter(str.isdigit, excel_cel)))
        row_index = row_number - 1
        col_index = ord(col_letter.upper()) - ord('A')

        waarde = df.iat[row_index, col_index]
        
        if pd.isna(waarde) or waarde == 0 or waarde == 0.0:
            return "DELETE"
        return str(waarde)
    except (IndexError, ValueError):
        return "DELETE"

def replace_placeholder_text(element, vervangingen):
    for placeholder, waarde in vervangingen.items():
        pattern = rf'\b{re.escape(placeholder)}\b'
        element.text = re.sub(pattern, waarde, element.text)

def verwijder_delete_paragrafen(doc):
    for para in doc.paragraphs:
        while "DELETE" in para.text:
            text = para.text
            
            # Zoek alle DELETE-locaties
            delete_matches = list(re.finditer(r'DELETE', text))
            if not delete_matches:
                print("Geen DELETE's gevonden, stoppen.")
                break
            
            # Pak de eerste DELETE die we gaan verwerken
            delete_index = delete_matches[0].start()
            print(f"\nDELETE gevonden op index {delete_index} in tekst: {text}")
            
            # Zoek links de laatst voorkomende "Factuur" of "Ontvangst"
            left_match = list(re.finditer(r'(Factuur|Ontvangst)', text[:delete_index], re.IGNORECASE))
            if left_match:
                start_index = left_match[-1].start()
                print(f"Laatste 'Factuur' of 'Ontvangst' gevonden op index {start_index}")
            else:
                print("Geen 'Factuur' of 'Ontvangst' gevonden voor deze DELETE. Stop met verwerken.")
                break  # Geen juiste match, stop verwerking
            
            # Zoek rechts de eerstvolgende DELETE
            right_match = next((m for m in delete_matches if m.start() > delete_index), None)
            if right_match:
                end_index = right_match.end()
                print(f"Volgende DELETE gevonden op index {end_index}")
            else:
                end_index = delete_index + len("DELETE")
                print("Geen extra DELETE gevonden, nemen einde van huidige DELETE.")
            
            # Toon het exacte deel dat wordt verwijderd
            print(f"Verwijderen: '{text[start_index:end_index]}'")
            
            # Verwijder het ongewenste deel en behoud de rest
            nieuwe_text = text[:start_index].strip() + "\n" + text[end_index:].strip()
            print(f"Overgebleven tekst na verwijdering: {nieuwe_text}")
            
            # Voeg extra nieuwe regels toe waar nodig voor correcte formatting
            nieuwe_text = re.sub(r'(Buitengerechtelijke incassokosten|Subtotaal)', r'\n\1', nieuwe_text)
            
            # Update de paragraaftekst
            para.text = nieuwe_text

def verwijder_onnodige_spaties(doc):
    """
    Verwijdert onnodige lege regels en overtollige witruimte in het document,
    terwijl de oorspronkelijke opmaak (bold, italic) per teken behouden blijft.
    """
    for para in doc.paragraphs:
        # Check of de paragraaf runs heeft
        if not para.runs:
            continue
        
        # Verzamel alle runs met hun opmaak
        runs_with_format = []
        for run in para.runs:
            runs_with_format.append({
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic
            })
        
        # Combineer runs om te kunnen bewerken
        full_text = "".join(run['text'] for run in runs_with_format)
        
        # Verwijder extra witregels en spaties
        cleaned_text = re.sub(r'\n{3,}', '\n\n', full_text).strip()
        cleaned_text = re.sub(r' {2,}', ' ', cleaned_text)  # Meerdere spaties naar één spatie
        
        # Als er geen wijzigingen zijn, sla deze paragraaf over
        if cleaned_text == full_text:
            continue
        
        # We moeten de opmaak van de oorspronkelijke tekst behouden
        # Eerst leegmaken
        para.clear()
        
        # Als alle tekst is verwijderd, ga door naar de volgende paragraaf
        if not cleaned_text:
            continue
            
        # Voeg één nieuwe run toe met dezelfde opmaak als de oorspronkelijke runs
        # Hier gebruiken we de dominante opmaak (meest voorkomend in de paragraaf)
        bold_count = sum(1 for run in runs_with_format if run['bold'])
        italic_count = sum(1 for run in runs_with_format if run['italic'])
        
        # Bepaal of de meerderheid bold/italic was
        was_bold = bold_count > len(runs_with_format) / 2
        was_italic = italic_count > len(runs_with_format) / 2
        
        # Voeg de opgeschoonde tekst toe met de juiste opmaak
        new_run = para.add_run(cleaned_text)
        new_run.bold = was_bold
        new_run.italic = was_italic

def process_document(excel_path, word_path, output_path):
    try:
        # Laad Excel
        df_excel = pd.read_excel(excel_path, header=None)
        
        # Maak vervangingen dictionary
        vervangingen = {placeholder: haal_excel_waarde(df_excel, cel) for placeholder, cel in mapping.items()}
        
        # Open Word document
        doc = Document(word_path)
        
        # Vervang placeholders
        for para in doc.paragraphs:
            for run in para.runs:
                replace_placeholder_text(run, vervangingen)
        
        # Verwijder DELETE paragrafen
        verwijder_delete_paragrafen(doc)
        
        # Verwijder onnodige spaties
        verwijder_onnodige_spaties(doc)
        
        # Opslaan naar opgegeven output pad
        doc.save(output_path)
        
        print(f"Document met correcte vervangingen en opgeschoonde 'DELETE'-secties opgeslagen als: {output_path}")
        return True
    except Exception as e:
        print(f"Fout bij verwerken document: {str(e)}")
        return False

if __name__ == '__main__':
    # Bij aanroep vanaf command line (via deel_3.py):
    # python main.py excel_path word_path output_path
    if len(sys.argv) == 4:
        excel_path = sys.argv[1]
        word_path = sys.argv[2]
        output_path = sys.argv[3]
        success = process_document(excel_path, word_path, output_path)
        # Zorg ervoor dat het script een correcte returncode teruggeeft
        if not success:
            sys.exit(1)